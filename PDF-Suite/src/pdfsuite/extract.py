"""
PT-PT: Leitura de documentos.

       Um formato, uma funcao, e uma unica saida: um `Documento` com o texto
       ja extraido. Quem compara ou resume nao precisa de saber se aquilo veio
       de um PDF ou de um Word.

       A distincao que este modulo faz questao de nao perder e entre um
       documento vazio e um PDF digitalizado. Sao os dois «sem texto», mas o
       primeiro nao tem nada la dentro e o segundo tem tudo — em imagem. Dizer
       ao utilizador «o ficheiro esta vazio» quando o problema e falta de OCR
       manda-o procurar no sitio errado.

EN-UK: Document reading.

       One format, one function, one output: a `Documento` with the text
       already extracted. Whoever compares or summarises need not know whether
       it came from a PDF or a Word file.

       The distinction this module takes care not to lose is between an empty
       document and a scanned PDF. Both are "no text", but the first has
       nothing inside and the second has everything — as an image.

Created by Redfox using Claude
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from .models import Documento

log = logging.getLogger(__name__)

# PT-PT: Formatos que sabemos ler. A interface usa isto para o filtro do
#        selector de ficheiros, para nao haver duas listas a divergir.
# EN-UK: Formats we can read. The interface uses this for the file dialog
#        filter, so there are not two lists drifting apart.
EXTENSOES_PDF: tuple[str, ...] = (".pdf",)
EXTENSOES_WORD: tuple[str, ...] = (".docx",)
EXTENSOES_TEXTO: tuple[str, ...] = (".txt", ".md", ".csv", ".log", ".rtf", ".json")
EXTENSOES = EXTENSOES_PDF + EXTENSOES_WORD + EXTENSOES_TEXTO

# PT-PT: Abaixo disto por pagina, um PDF e quase de certeza digitalizado. O
#        valor nao e zero de proposito: paginas digitalizadas trazem muitas
#        vezes um cabecalho ou um numero de pagina em texto real, vindos do
#        proprio scanner. Zero como limite deixava-os passar por bons.
# EN-UK: Below this per page, a PDF is almost certainly scanned. The value is
#        deliberately not zero: scanned pages often carry a header or page
#        number as real text, produced by the scanner itself.
CARACTERES_MINIMOS_POR_PAGINA = 60


def formatos_suportados() -> str:
    """PT-PT: Lista legivel para mensagens. / EN-UK: Readable list for messages."""
    return ", ".join(sorted(EXTENSOES))


def _limpar(texto: str) -> str:
    """
    PT-PT: Normaliza o texto extraido.

           A juncao de palavras cortadas no fim da linha («fornece-\ndor») nao
           e cosmetica: sem ela, procurar «fornecedor» no texto de um PDF
           justificado falha exactamente nas paginas onde a palavra e mais
           provavel de aparecer.

    EN-UK: Normalises extracted text. Rejoining words hyphenated across line
           breaks is not cosmetic: without it, searching for a word in a
           justified PDF fails precisely on the pages where it is most likely
           to appear.
    """
    if not texto:
        return ""

    texto = texto.replace("\u00ad", "")  # hífen condicional
    texto = re.sub(r"(\w)-\n(\w)", r"\1\2", texto)
    texto = texto.replace("\r\n", "\n").replace("\r", "\n")
    texto = re.sub(r"[\u00a0\u202f\u2009]", " ", texto)
    texto = re.sub(r"[ \t]+", " ", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()


def ler_pdf(caminho: Path) -> Documento:
    """
    PT-PT: Le um PDF, com duas estrategias em cascata.

           Primeiro o pdfplumber, que respeita a disposicao das colunas — numa
           tabela de precos, e a diferenca entre ler «Artigo A 250,00» e ler a
           coluna dos artigos toda seguida da coluna dos precos toda. Se
           falhar, o pypdf, que e mais tolerante a PDF mal formados.

    EN-UK: Reads a PDF with two strategies in cascade. First pdfplumber, which
           respects column layout — on a price table, that is the difference
           between reading "Item A 250.00" and reading the whole item column
           followed by the whole price column. If that fails, pypdf, which is
           more tolerant of malformed PDFs.
    """
    documento = Documento(caminho=caminho, formato="PDF")

    try:
        import pdfplumber

        with pdfplumber.open(str(caminho)) as pdf:
            documento.paginas = len(pdf.pages)
            partes = []
            for pagina in pdf.pages:
                try:
                    partes.append(pagina.extract_text() or "")
                except Exception as exc:  # noqa: BLE001
                    # PT-PT: Uma pagina ilegivel nao pode custar o documento
                    #        inteiro. Numa proposta de 40 paginas, perder a
                    #        pagina 12 e mau; perder as 40 e inutilizavel.
                    # EN-UK: One unreadable page must not cost the whole
                    #        document.
                    log.warning("Página ilegível em %s: %s", caminho.name, exc)
                    partes.append("")
            documento.texto = _limpar("\n\n".join(partes))
    except Exception as exc:  # noqa: BLE001
        log.warning("pdfplumber falhou em %s (%s); a tentar o pypdf.", caminho.name, exc)
        try:
            from pypdf import PdfReader

            leitor = PdfReader(str(caminho))
            documento.paginas = len(leitor.pages)
            documento.texto = _limpar(
                "\n\n".join((p.extract_text() or "") for p in leitor.pages)
            )
        except Exception as exc2:  # noqa: BLE001
            documento.erro = f"Não foi possível ler o PDF: {exc2}"
            log.error("Falha ao ler %s: %s", caminho, exc2)
            return documento

    if documento.paginas and len(documento.texto) < CARACTERES_MINIMOS_POR_PAGINA * documento.paginas:
        documento.digitalizado = True
        if not documento.texto.strip():
            documento.erro = (
                "PDF sem camada de texto — está digitalizado como imagem. "
                "É preciso passar por OCR antes de poder ser analisado."
            )

    return documento


def ler_docx(caminho: Path) -> Documento:
    """
    PT-PT: Le um documento Word.

           Le tambem as tabelas, e nao so os paragrafos. Numa proposta
           comercial, os precos estao quase sempre numa tabela: ignora-las
           deixava o documento com o texto de cortesia e sem um unico numero.

    EN-UK: Reads a Word document, including tables and not just paragraphs. In
           a commercial quote the prices are almost always in a table; ignoring
           them left the document with the courtesy text and not a single
           figure.
    """
    documento = Documento(caminho=caminho, formato="Word")

    try:
        import docx

        ficheiro = docx.Document(str(caminho))
    except Exception as exc:  # noqa: BLE001
        documento.erro = f"Não foi possível ler o documento Word: {exc}"
        log.error("Falha ao ler %s: %s", caminho, exc)
        return documento

    partes = [p.text for p in ficheiro.paragraphs if p.text.strip()]

    for tabela in ficheiro.tables:
        for linha in tabela.rows:
            celulas = [c.text.strip() for c in linha.cells]
            if any(celulas):
                # PT-PT: As celulas sao separadas por tabulacao para a linha
                #        continuar a ler-se como uma linha de tabela e nao como
                #        prosa colada.
                # EN-UK: Cells joined by tabs so the row still reads as a table
                #        row rather than run-together prose.
                partes.append("\t".join(celulas))

    documento.texto = _limpar("\n".join(partes))
    documento.paginas = 0
    return documento


def ler_texto(caminho: Path) -> Documento:
    """
    PT-PT: Le um ficheiro de texto.

           Tenta UTF-8 e depois cp1252. E a ordem certa para Portugal: os
           ficheiros novos sao UTF-8, e os antigos, exportados por aplicacoes
           de gestao em Windows, sao cp1252. Ler cp1252 como UTF-8 rebenta;
           ler UTF-8 como cp1252 nao rebenta, mas estraga todos os acentos em
           silencio, que e pior. Por isso o UTF-8 vem primeiro.

    EN-UK: Reads a text file. Tries UTF-8 then cp1252 — the right order:
           reading cp1252 as UTF-8 raises, while reading UTF-8 as cp1252 does
           not raise but silently ruins every accent, which is worse.
    """
    documento = Documento(caminho=caminho, formato=caminho.suffix.lstrip(".").upper())

    for codificacao in ("utf-8", "cp1252", "latin-1"):
        try:
            bruto = caminho.read_text(encoding=codificacao)
            break
        except UnicodeDecodeError:
            continue
        except OSError as exc:
            documento.erro = f"Não foi possível abrir o ficheiro: {exc}"
            return documento
    else:
        documento.erro = "Não foi possível determinar a codificação do ficheiro."
        return documento

    if caminho.suffix.lower() == ".rtf":
        bruto = _rtf_para_texto(bruto)

    documento.texto = _limpar(bruto)
    return documento


def _rtf_para_texto(bruto: str) -> str:
    """
    PT-PT: Extrai o texto de um RTF sem bibliotecas externas.

           E uma aproximacao, nao um interpretador de RTF: tira os grupos de
           controlo e converte os escapes hexadecimais. Chega para um documento
           de texto corrido exportado do Word, que e o caso em que estes
           ficheiros ainda aparecem, e nao chega para um RTF com tabelas ou
           imagens — o utilizador e avisado disso na interface.

    EN-UK: Extracts text from RTF without external libraries. An approximation,
           not an RTF interpreter: enough for a plain document exported from
           Word, not enough for one with tables or images.
    """
    texto = re.sub(r"\{\\\*.*?\}", "", bruto, flags=re.DOTALL)
    texto = re.sub(r"\\'([0-9a-fA-F]{2})", lambda m: chr(int(m.group(1), 16)), texto)
    texto = re.sub(r"\\par[d]?\b", "\n", texto)
    texto = re.sub(r"\\[a-zA-Z]+-?\d*\s?", "", texto)
    texto = texto.replace("{", "").replace("}", "")
    return texto


def ler(caminho: Path | str) -> Documento:
    """
    PT-PT: Le um documento, escolhendo a estrategia pela extensao.
    EN-UK: Reads a document, picking the strategy from the extension.
    """
    caminho = Path(caminho)

    if not caminho.exists():
        return Documento(caminho=caminho, erro="O ficheiro não existe.")
    if not caminho.is_file():
        return Documento(caminho=caminho, erro="O caminho não é um ficheiro.")

    try:
        if caminho.stat().st_size == 0:
            return Documento(caminho=caminho, erro="O ficheiro está vazio (0 bytes).")
    except OSError as exc:
        return Documento(caminho=caminho, erro=f"Não foi possível aceder ao ficheiro: {exc}")

    sufixo = caminho.suffix.lower()
    if sufixo in EXTENSOES_PDF:
        return ler_pdf(caminho)
    if sufixo in EXTENSOES_WORD:
        return ler_docx(caminho)
    if sufixo in EXTENSOES_TEXTO:
        return ler_texto(caminho)

    if sufixo == ".doc":
        # PT-PT: O .doc antigo e um formato binario diferente do .docx e nao ha
        #        forma de o ler sem o Word ou o LibreOffice instalados. Dizer
        #        isto e o que fazer a seguir vale mais do que um erro generico.
        # EN-UK: The old .doc is a different binary format and cannot be read
        #        without Word or LibreOffice installed.
        return Documento(
            caminho=caminho,
            erro=(
                "Formato .doc antigo não suportado. Abra no Word e grave como "
                ".docx ou .pdf — são dois cliques e o resultado é melhor do que "
                "qualquer conversão automática."
            ),
        )

    return Documento(
        caminho=caminho,
        erro=f"Formato não suportado ({sufixo or 'sem extensão'}). Aceites: {formatos_suportados()}.",
    )


def ler_varios(caminhos: list[Path | str]) -> list[Documento]:
    """
    PT-PT: Le varios documentos. Os que falharem vem na lista com o erro
           preenchido, em vez de desaparecerem — numa comparacao de seis
           propostas, ficar com cinco sem perceber qual faltou e pior do que
           nao ter nenhuma.
    EN-UK: Reads several documents. Failures come back in the list with the
           error filled in rather than vanishing: in a six-way comparison,
           ending up with five and not knowing which one dropped out is worse
           than having none.
    """
    return [ler(c) for c in caminhos]
